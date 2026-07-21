"""UTF-8 text and JSON to Word conversion using python-docx."""

from __future__ import annotations

import io
import json
from pathlib import Path

# Large bodies: one giant ``add_run`` is very slow for multi‑MB strings; several runs in
# one paragraph concatenate to the same ``paragraph.text`` for extraction.
_RUN_CHUNK_CHARS = 262_144


def json_pretty_or_raw(raw: str) -> str:
    """Return ``json.dumps(..., indent=2)`` when ``raw`` parses; else ``raw`` unchanged."""
    if not raw.strip():
        return ""
    try:
        obj = json.loads(raw)
        return json.dumps(obj, ensure_ascii=False, indent=2)
    except Exception:
        return raw


def plain_text_to_docx_bytes(
    text: str,
    *,
    title: str | None = None,
    fallback_title: str = "document",
) -> io.BytesIO:
    """Write ``text`` to a minimal ``.docx`` (single body paragraph, newlines preserved)."""
    from docx import Document

    doc = Document()
    doc.core_properties.title = (title or "").strip() or fallback_title
    para = doc.add_paragraph()
    n = len(text)
    if n <= _RUN_CHUNK_CHARS:
        para.add_run(text)
    else:
        i = 0
        while i < n:
            para.add_run(text[i : i + _RUN_CHUNK_CHARS])
            i += _RUN_CHUNK_CHARS
    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf


def txt_path_to_docx_bytes(
    path: Path,
    *,
    title: str | None = None,
    encoding: str = "utf-8",
    errors: str = "replace",
) -> io.BytesIO:
    """Read a text file and return ``.docx`` bytes."""
    raw = path.read_text(encoding=encoding, errors=errors)
    return plain_text_to_docx_bytes(raw, title=title, fallback_title=path.stem)


def json_string_to_docx_bytes(
    raw_json: str,
    *,
    title: str | None = None,
    fallback_title: str = "document",
) -> io.BytesIO:
    """Pretty-print JSON when valid; otherwise embed the raw string."""
    body = json_pretty_or_raw(raw_json)
    return plain_text_to_docx_bytes(body, title=title, fallback_title=fallback_title)


def json_path_to_docx_bytes(
    path: Path,
    *,
    title: str | None = None,
    encoding: str = "utf-8",
    errors: str = "replace",
) -> io.BytesIO:
    """Read ``.json`` (or any UTF-8 file treated as JSON) and return ``.docx`` bytes."""
    raw = path.read_text(encoding=encoding, errors=errors)
    return json_string_to_docx_bytes(raw, title=title, fallback_title=path.stem)
